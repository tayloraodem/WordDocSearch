import docx
import os

def search (keyword):
    print("Searching for: " + keyword + "\n~~~~~~~~~~~~~~~~~~~~~~~~~~")

    directory = os.listdir()[:23]
    print(directory)
    
    chapters = [directory[0]] + directory[12:] + directory[1:12]
    
    for file in chapters:
        print("\n" + file)
        doc = docx.Document(file)
        result = [p.text for p in doc.paragraphs]
        p = 1
        count = 0
        for para in result:
            num = para.count(keyword)
            if num > 0:
                print("Found " + str(num) + " instance(s) in paragraph: " + str(p))
                count = count + num
            p = p + 1
        print("Total: " + str(count))


def count():

    directory = os.listdir()[:25]
    
    
    chapters = [directory[0]] + [directory[11]] + directory[16:] + directory[1:10]
    print(chapters)
    
    for file in chapters:
        print("\n" + file)
        doc = docx.Document(file)
        result = [p.text for p in doc.paragraphs]
        p = 1
        count = 0
        for para in result:
            para = para.split(" ")

            
            num = len(para)
            count = count + num

        print("Total words: " + str(count))
        die
